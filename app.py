import os
from pathlib import Path
from flask import Flask, render_template_string, request, jsonify, send_file
from pypdf import PdfReader
from openai import OpenAI, AuthenticationError, BadRequestError, APIStatusError
from dotenv import load_dotenv
from datetime import datetime
from io import BytesIO

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

BASE_DIR = Path(__file__).resolve().parent
load_dotenv(dotenv_path=BASE_DIR / ".env", override=True)

app = Flask(__name__)

# Store last analysis result for export
last_analysis = {
    "resume_text": "",
    "job_description": "",
    "analysis": "",
    "timestamp": None,
}

def _clean_env_key(value: str) -> str:
    return (value or "").strip().strip('"').strip("'")


def _is_real_api_key(value: str) -> bool:
    if not value:
        return False
    lowered = value.lower()
    placeholders = (
        "your-key-here",
        "your-openai-key-here",
        "your-perplexity-key-here",
        "replace-me",
        "changeme",
        "dummy",
        "test",
    )
    return not any(token in lowered for token in placeholders)


pplx_api_key = _clean_env_key(os.getenv("PPLX_API_KEY"))
openai_api_key = _clean_env_key(os.getenv("OPENAI_API_KEY"))

def _build_clients():
    available = []
    if _is_real_api_key(pplx_api_key):
        available.append(("perplexity", OpenAI(api_key=pplx_api_key, base_url="https://api.perplexity.ai")))
    if _is_real_api_key(openai_api_key):
        available.append(("openai", OpenAI(api_key=openai_api_key)))
    return available


clients = _build_clients()
provider = clients[0][0] if clients else None
client = clients[0][1] if clients else None

def extract_resume_text(file_storage):
    filename = (file_storage.filename or "").lower()
    if filename.endswith(".pdf"):
        reader = PdfReader(file_storage)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or "")
        return "\n".join(text).strip()
    elif filename.endswith(".txt"):
        return file_storage.read().decode("utf-8", errors="ignore").strip()
    else:
        return ""

@app.route("/", methods=["GET"])
def home():
    html_path = BASE_DIR / "index.html"
    if html_path.exists():
        with open(html_path, "r", encoding="utf-8") as f:
            return f.read()
    return "<h1>AI Resume Scanner</h1><p>index.html not found.</p>", 404

@app.route("/favicon.ico", methods=["GET"])
def favicon():
        return "", 204

@app.route("/scan", methods=["POST"])
def scan():
    try:
        if not client:
            return jsonify({
                "error": "API key not configured. Add a valid PPLX_API_KEY or OPENAI_API_KEY in .env (same folder as app.py), then fully restart the app."
            }), 500

        resume_file = request.files.get("resume")
        job_description = (request.form.get("job_description") or "").strip()

        if not resume_file or not job_description:
            return jsonify({"error": "Resume file and job description are required."}), 400

        resume_text = extract_resume_text(resume_file)
        if not resume_text:
            return jsonify({"error": "Could not extract resume text. Use a readable PDF or TXT."}), 400

        prompt = f"""
You are an AI resume coach and hiring expert. Provide human-like, actionable guidance.

Compare this resume with the job description and return:
1) Match score 0-100 (brief rationale)
2) Core strengths aligned to the role (bullets)
3) Matched keywords and related terms found in the resume (bullets)
4) Missing keywords (bullets)
5) Related terms/synonyms to consider (bullets)
6) Improvement suggestions (bullets)
7) Rewritten professional summary (3-4 lines)

Resume:
{resume_text}

Job Description:
{job_description}
"""

        analysis_text = ""
        auth_failures = 0

        for active_provider, active_client in clients:
            try:
                if active_provider == "perplexity":
                    response = active_client.chat.completions.create(
                        model="sonar-pro",
                        messages=[
                            {"role": "system", "content": "You are an AI resume coach and hiring expert. Provide human-like, actionable guidance."},
                            {"role": "user", "content": prompt},
                        ],
                    )
                    analysis_text = (response.choices[0].message.content or "").strip()
                else:
                    response = active_client.responses.create(
                        model="gpt-4.1-mini",
                        input=prompt
                    )
                    analysis_text = (response.output_text or "").strip()

                if analysis_text:
                    break
            except AuthenticationError:
                auth_failures += 1
                continue

        if auth_failures == len(clients):
            return jsonify({
                "error": "Invalid API key (401). Update PPLX_API_KEY or OPENAI_API_KEY in .env and restart the app."
            }), 401

        if not analysis_text:
            return jsonify({"error": "Model returned an empty response. Please try again."}), 502

        # Store results for export
        global last_analysis
        last_analysis = {
            "resume_text": resume_text,
            "job_description": job_description,
            "analysis": analysis_text,
            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        }

        return jsonify({"analysis": analysis_text})

    except BadRequestError as e:
        return jsonify({"error": f"OpenAI request error: {e}"}), 400
    except APIStatusError as e:
        return jsonify({"error": f"OpenAI API status error ({e.status_code}): {e}"}), 502

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/export-excel", methods=["GET"])
def export_excel():
    if not last_analysis.get("analysis"):
        return jsonify({"error": "No analysis to export. Run a scan first."}), 400

    try:
        def _pdf_escape(text):
            return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

        def _build_pdf(text_lines):
            lines_per_page = 45
            page_height = 792
            top_margin = 760
            line_height = 15

            pages = [text_lines[i:i + lines_per_page] for i in range(0, len(text_lines), lines_per_page)]
            if not pages:
                pages = [["AI Resume Scanner - Analysis Report"]]

            objects = []
            object_ids = {}
            next_id = 1

            def new_object(content):
                nonlocal next_id
                object_ids[next_id] = content
                next_id += 1
                return next_id - 1

            font_id = new_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
            page_ids = []
            content_ids = []

            pages_root_id = next_id
            object_ids[pages_root_id] = ""
            next_id += 1

            for page_lines in pages:
                stream_lines = ["BT", "/F1 11 Tf"]
                y = top_margin
                for line in page_lines:
                    safe = _pdf_escape(line)
                    stream_lines.append(f"1 0 0 1 40 {y} Tm ({safe}) Tj")
                    y -= line_height
                stream_lines.append("ET")
                stream_data = "\n".join(stream_lines).encode("latin-1", errors="replace")
                content_id = new_object(f"<< /Length {len(stream_data)} >>\nstream\n" + stream_data.decode("latin-1") + "\nendstream")
                content_ids.append(content_id)

                page_id = new_object(f"<< /Type /Page /Parent {pages_root_id} 0 R /MediaBox [0 0 612 {page_height}] /Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {content_id} 0 R >>")
                page_ids.append(page_id)

            object_ids[pages_root_id] = f"<< /Type /Pages /Count {len(page_ids)} /Kids [{' '.join(f'{pid} 0 R' for pid in page_ids)}] >>"
            catalog_id = new_object(f"<< /Type /Catalog /Pages {pages_root_id} 0 R >>")

            pdf = BytesIO()
            pdf.write(b"%PDF-1.4\n")

            offsets = {}
            for obj_id in sorted(object_ids.keys()):
                offsets[obj_id] = pdf.tell()
                pdf.write(f"{obj_id} 0 obj\n".encode("ascii"))
                pdf.write((object_ids[obj_id] + "\n").encode("latin-1", errors="replace"))
                pdf.write(b"endobj\n")

            xref_start = pdf.tell()
            total_objects = max(object_ids.keys())
            pdf.write(f"xref\n0 {total_objects + 1}\n".encode("ascii"))
            pdf.write(b"0000000000 65535 f \n")
            for obj_id in range(1, total_objects + 1):
                offset = offsets.get(obj_id, 0)
                pdf.write(f"{offset:010} 00000 n \n".encode("ascii"))

            pdf.write(f"trailer\n<< /Size {total_objects + 1} /Root {catalog_id} 0 R >>\nstartxref\n{xref_start}\n%%EOF".encode("ascii"))
            pdf.seek(0)
            return pdf

        report_lines = [
            "AI Resume Scanner - Analysis Report",
            f"Generated: {last_analysis.get('timestamp', 'N/A')}",
            "",
            "Resume Text",
            *last_analysis.get("resume_text", "").split("\n"),
            "",
            "Job Description",
            *last_analysis.get("job_description", "").split("\n"),
            "",
            "AI Analysis & Recommendations",
            *last_analysis.get("analysis", "").split("\n"),
        ]

        output = _build_pdf(report_lines)

        timestamp = last_analysis.get("timestamp", "").replace(" ", "_").replace(":", "-")
        filename = f"resume_analysis_{timestamp}.pdf"

        return send_file(
            output,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({"error": f"Export failed: {str(e)}"}), 500

@app.route("/export-docx", methods=["GET"])
def export_docx():
    if not DOCX_AVAILABLE:
        return jsonify({"error": "DOCX export not available. Install python-docx: pip install python-docx"}), 500

    if not last_analysis.get("analysis"):
        return jsonify({"error": "No analysis to export. Run a scan first."}), 400

    try:
        doc = Document()
        title = doc.add_heading("AI Resume Scanner - Analysis Report", level=1)
        for run in title.runs:
            run.font.size = Pt(16)

        doc.add_paragraph(f"Generated: {last_analysis.get('timestamp', 'N/A')}")

        doc.add_heading("Resume Text", level=2)
        doc.add_paragraph(last_analysis.get("resume_text", ""))

        doc.add_heading("Job Description", level=2)
        doc.add_paragraph(last_analysis.get("job_description", ""))

        doc.add_heading("AI Analysis & Recommendations", level=2)
        doc.add_paragraph(last_analysis.get("analysis", ""))

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        timestamp = last_analysis.get("timestamp", "").replace(" ", "_").replace(":", "-")
        filename = f"resume_analysis_{timestamp}.docx"

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({"error": f"DOCX export failed: {str(e)}"}), 500

if __name__ == "__main__":
    app.run(debug=True)